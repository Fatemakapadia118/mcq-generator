import random
import os
from typing import List, Dict

import nltk
from nltk.corpus import stopwords, wordnet as wn
from nltk.tokenize import sent_tokenize, word_tokenize

import pdfplumber
from docx import Document

# Ensure required NLTK resources are available
def ensure_nltk_resources():
    resources = [
        "punkt",
        "punkt_tab",
        "wordnet",
        "omw-1.4",
        "averaged_perceptron_tagger",
        "averaged_perceptron_tagger_eng",  # NEW FIX
        "stopwords"
    ]
    for res in resources:
        try:
            if res == "punkt":
                nltk.data.find(f"tokenizers/{res}")
            elif res == "averaged_perceptron_tagger_eng":
                nltk.data.find("taggers/averaged_perceptron_tagger_eng")
            elif res == "averaged_perceptron_tagger":
                nltk.data.find("taggers/averaged_perceptron_tagger")
            else:
                nltk.data.find(res)
        except LookupError:
            nltk.download(res)

ensure_nltk_resources()

# Additional downloads (quiet mode to suppress logs)
nltk.download("punkt", quiet=True)
nltk.download("punkt_tab", quiet=True)
nltk.download("averaged_perceptron_tagger", quiet=True)
nltk.download("wordnet", quiet=True)
nltk.download("omw-1.4", quiet=True)
nltk.download("stopwords", quiet=True)

STOPWORDS = set(stopwords.words("english"))


def extract_text_from_pdf(path: str) -> str:
    import pdfplumber
    text_parts = []
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)
    return "\n".join(text_parts)


def extract_text_from_docx(path: str) -> str:
    from docx import Document
    doc = Document(path)
    return "\n".join([p.text for p in doc.paragraphs if p.text.strip()])


def _select_keyword(sentence: str, difficulty="medium") -> str:
    tokens = word_tokenize(sentence)
    tags = nltk.pos_tag(tokens)
    nouns = [w for w, t in tags if t.startswith("NN") and w.lower() not in STOPWORDS]

    if not nouns:
        return ""

    if difficulty == "easy":
        return min(nouns, key=len)  # shorter, simpler
    elif difficulty == "hard":
        return max(nouns, key=len)  # longer, technical
    return random.choice(nouns)  # medium → random noun


def _get_distractors(answer: str, max_d=3) -> List[str]:
    distractors = set()
    for syn in wn.synsets(answer):
        for lemma in syn.lemmas():
            w = lemma.name().replace("_", " ")
            if w.lower() != answer.lower():
                distractors.add(w)
    return list(distractors)[:max_d]


def generate_mcqs_from_text(text: str, num_questions: int = 5, difficulty="medium") -> List[Dict]:
    sentences = [s for s in sent_tokenize(text) if len(s.split()) > 5]
    random.shuffle(sentences)

    mcqs = []
    for sent in sentences[: num_questions * 2]:
        keyword = _select_keyword(sent, difficulty)
        if not keyword:
            continue
        question = sent.replace(keyword, "_____")
        correct = keyword
        distractors = _get_distractors(correct)

        while len(distractors) < 3:
            distractors.append(correct + random.choice(["s", "ing", "ed"]))
        options = distractors[:3] + [correct]
        random.shuffle(options)

        mcqs.append({"question": question, "options": options, "answer": correct})

        if len(mcqs) >= num_questions:
            break

    return mcqs


def export_docx(mcqs: List[Dict], out_path: str):
    from docx import Document
    doc = Document()
    doc.add_heading("Generated MCQs", 0)
    for i, q in enumerate(mcqs, 1):
        doc.add_paragraph(f"Q{i}. {q['question']}")
        for j, opt in enumerate(q["options"]):
            doc.add_paragraph(f"   {chr(65+j)}) {opt}")
    doc.save(out_path)
